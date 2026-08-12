"""Document -> text extraction.

Primary backend is `docling` (rich layout-aware parsing). Because docling pulls in a
large ML stack that may still be installing, every path degrades gracefully to a
lightweight per-format fallback (pypdf / python-docx / openpyxl / plain text) so the
core pipeline is runnable immediately.
"""
from __future__ import annotations

import logging
from dataclasses import dataclass
from pathlib import Path
from typing import Any

logger = logging.getLogger(__name__)

SUPPORTED_EXTS = {"pdf", "docx", "doc", "xlsx", "xls", "txt", "md", "markdown", "csv"}
_DOCLING_EXTS = {"pdf", "docx", "xlsx", "md", "markdown", "html", "pptx"}

# Cache the docling converter (construction is expensive).
_docling_converter = None
_docling_unavailable = False


@dataclass
class ParseResult:
    text: str
    backend: str  # "docling" | "fallback:pdf" | ...
    structured_document: Any | None = None


def _get_docling():
    global _docling_converter, _docling_unavailable
    if _docling_unavailable:
        return None
    if _docling_converter is None:
        try:
            from docling.document_converter import DocumentConverter

            _docling_converter = DocumentConverter()
        except Exception as e:  # noqa: BLE001  (not installed / import failure)
            logger.info("docling unavailable, using fallback parsers: %s", e)
            _docling_unavailable = True
            return None
    return _docling_converter


def _try_docling(path: Path, ext: str) -> ParseResult | None:
    if ext not in _DOCLING_EXTS:
        return None
    conv = _get_docling()
    if conv is None:
        return None
    try:
        result = conv.convert(str(path))
        return ParseResult(
            text=result.document.export_to_markdown(),
            backend="docling",
            structured_document=result.document,
        )
    except Exception as e:  # noqa: BLE001
        logger.warning("docling failed on %s (%s); falling back", path.name, e)
        return None


# --------------------------------------------------------------------------- #
# Lightweight fallbacks
# --------------------------------------------------------------------------- #
def _fallback_pdf(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    parts = []
    for i, page in enumerate(reader.pages):
        parts.append(f"## Page {i + 1}\n\n{page.extract_text() or ''}")
    return "\n\n".join(parts)


def _fallback_docx(path: Path) -> str:
    import docx

    d = docx.Document(str(path))
    parts = [p.text for p in d.paragraphs if p.text.strip()]
    for table in d.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n\n".join(parts)


def _fallback_xlsx(path: Path) -> str:
    from openpyxl import load_workbook

    wb = load_workbook(str(path), read_only=True, data_only=True)
    parts = []
    for ws in wb.worksheets:
        parts.append(f"## Sheet: {ws.title}")
        for row in ws.iter_rows(values_only=True):
            cells = ["" if v is None else str(v) for v in row]
            if any(c.strip() for c in cells):
                parts.append("\t".join(cells))
    return "\n".join(parts)


def _fallback_text(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


def _fallback_parse(path: Path, ext: str) -> ParseResult:
    if ext == "pdf":
        return ParseResult(_fallback_pdf(path), "fallback:pdf")
    if ext in ("docx", "doc"):
        return ParseResult(_fallback_docx(path), "fallback:docx")
    if ext in ("xlsx", "xls"):
        return ParseResult(_fallback_xlsx(path), "fallback:xlsx")
    if ext in ("txt", "md", "markdown", "csv"):
        return ParseResult(_fallback_text(path), "fallback:text")
    raise ValueError(f"Unsupported file type: .{ext}")


def parse_file(path: Path, ext: str) -> ParseResult:
    """Extract text from a file, preferring docling and degrading to fallbacks."""
    ext = ext.lower().lstrip(".")
    if ext not in SUPPORTED_EXTS:
        raise ValueError(f"Unsupported file type: .{ext}")

    result = _try_docling(path, ext)
    if result is not None and result.text.strip():
        return result
    return _fallback_parse(path, ext)
